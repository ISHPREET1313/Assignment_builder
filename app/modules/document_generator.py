"""
app/modules/document_generator.py  v4

Changes:
  - VS Code Dark+ syntax highlighting via Pygments (coloured Word runs)
  - Manual XML borders (no Table Grid dependency)
  - Dark code background (#1E1E1E) in the Word doc
"""

import datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Pygments ──────────────────────────────────────────────────────────────────
from pygments import lex
from pygments.lexers import get_lexer_by_name, guess_lexer_for_filename
from pygments.token import Token

# VS Code Dark+ palette  (token → hex string)
_VSCODE: dict = {
    "keyword"   : "569CD6",   # blue        — if, for, int, void …
    "keyword2"  : "C586C0",   # purple      — include, define
    "builtin"   : "DCDCAA",   # yellow      — printf, scanf, len …
    "string"    : "CE9178",   # orange-red  — "hello"
    "number"    : "B5CEA8",   # light green — 42, 3.14
    "comment"   : "6A9955",   # green       — // …
    "type"      : "4EC9B0",   # teal        — int, char, void (C types)
    "preproc"   : "C586C0",   # purple      — #include
    "operator"  : "D4D4D4",   # light grey  — + - * / = …
    "default"   : "D4D4D4",   # light grey  — everything else
    "bg"        : "1E1E1E",   # dark bg
}

def _vscode_color(ttype) -> str:
    """Map a Pygments token type to a VS Code Dark+ hex colour."""
    t = ttype
    while t:
        if   t in (Token.Keyword, Token.Keyword.Declaration,
                   Token.Keyword.Type, Token.Keyword.Reserved):
            return _VSCODE["keyword"]
        elif t in (Token.Keyword.Namespace,):           return _VSCODE["preproc"]
        elif t in (Token.Comment, Token.Comment.Single,
                   Token.Comment.Multiline):            return _VSCODE["comment"]
        elif str(t).startswith("Token.Literal.String"): return _VSCODE["string"]
        elif str(t).startswith("Token.Literal.Number"): return _VSCODE["number"]
        elif t in (Token.Name.Builtin,):                return _VSCODE["builtin"]
        elif t in (Token.Name.Function,):               return _VSCODE["builtin"]
        elif t in (Token.Comment.Preproc,
                   Token.Comment.PreprocFile):          return _VSCODE["preproc"]
        elif str(t).startswith("Token.Operator"):       return _VSCODE["operator"]
        t = t.parent
    return _VSCODE["default"]

def _hex_to_rgb(hex_: str) -> RGBColor:
    h = hex_.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


# ── XML helpers ───────────────────────────────────────────────────────────────

def _clear_body(doc):
    body   = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

def _add_borders(cell, color="444444", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcBorders to avoid duplicate XML (causes Word corruption)
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcB.append(el)
    tcPr.append(tcB)

def _shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd to avoid duplicate XML (causes Word corruption)
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def _section_label(doc, text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def _has_style(doc, name):
    return any(s.name == name for s in doc.styles)


# ── Syntax-highlighted code block ─────────────────────────────────────────────

def _code_block(doc, code: str, language: str):
    """
    Render code with VS Code Dark+ colours in a dark-background table cell.
    Each token gets its own run with the appropriate colour.
    """
    # Pick lexer
    lang_map = {"Python": "python", "C": "c", "C++": "cpp"}
    lex_name = lang_map.get(language, "text")
    try:
        lexer = get_lexer_by_name(lex_name, stripall=False)
    except Exception:
        lexer = get_lexer_by_name("text")

    tokens = list(lex(code, lexer))

    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _shade_cell(cell, _VSCODE["bg"])
    _add_borders(cell, color="333333", sz="4")

    # Language badge paragraph
    badge_para = cell.paragraphs[0]
    badge_para.paragraph_format.space_before = Pt(4)
    badge_para.paragraph_format.space_after  = Pt(0)
    br = badge_para.add_run(f" {language} ")
    br.font.name       = "Courier New"
    br.font.size       = Pt(8)
    br.bold            = True
    br.font.color.rgb  = RGBColor(0x00, 0x00, 0x00)
    # Badge background via highlight would need XML; just colour the text
    br.font.color.rgb  = _hex_to_rgb(_VSCODE["builtin"])

    # Code body — one paragraph per line, coloured runs per token
    current_para = None
    buffer_line  = []   # (text, color_hex) pairs for current line

    def _flush(para, segments):
        for text, color in segments:
            if not text:
                continue
            r = para.add_run(text)
            r.font.name      = "Courier New"
            r.font.size      = Pt(9)
            r.font.color.rgb = _hex_to_rgb(color)
        return []

    def _new_para():
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        return p

    current_para = _new_para()

    for ttype, value in tokens:
        color = _vscode_color(ttype)
        # Split on newlines — each \n starts a fresh paragraph
        parts = value.split("\n")
        for i, part in enumerate(parts):
            if i > 0:
                buffer_line = _flush(current_para, buffer_line)
                current_para = _new_para()
            if part:
                buffer_line.append((part, color))

    if buffer_line:
        _flush(current_para, buffer_line)

    doc.add_paragraph()


# ── Document generator ────────────────────────────────────────────────────────

class DocumentGenerator:
    def __init__(self, template_path: Path, out_dir: Path):
        self.template = Path(template_path)
        self.out_dir  = Path(out_dir)

    def generate(self, results: list) -> Path:
        doc = Document(str(self.template))
        _clear_body(doc)

        # Cover
        t = doc.add_heading("Programming Practicals", level=1)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if t.runs:
            t.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        sub = doc.add_paragraph(
            f"Automated Documentation Report  ·  "
            f"{datetime.datetime.now().strftime('%d %B %Y')}"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in sub.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        doc.add_paragraph()

        for entry in results:
            # exp_label is the display label in the doc; falls back to exp_no if not set
            label = entry.get("exp_label") or entry["exp_no"]
            h = doc.add_heading(f"Experiment {label}", level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            _section_label(doc, "Aim / Objective")
            aim = doc.add_paragraph(entry["question"])
            for r in aim.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(11)

            _section_label(doc, "Source Code")
            _code_block(doc, entry["source_code"], entry["language"])

            _section_label(doc, "Output")
            if entry["img_path"] and Path(entry["img_path"]).exists():
                doc.add_picture(str(entry["img_path"]), width=Inches(5.8))
            else:
                doc.add_paragraph("(output image not available)")

            doc.add_page_break()

        ts  = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        out = self.out_dir / f"PracticalReport_{ts}.docx"
        doc.save(str(out))
        return out
